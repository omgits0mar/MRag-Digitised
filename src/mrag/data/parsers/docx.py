"""DOCX parser using python-docx.

Emits a single ParsedDocument containing all paragraph text joined by
newlines. Tables are included cell-by-cell.
"""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path
from uuid import uuid4

from mrag.data.parsers.registry import ParsedDocument, ParserError


def parse_docx(path: Path, source_filename: str) -> Iterable[ParsedDocument]:
    """Extract paragraph and table-cell text from a .docx file."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dep is pinned
        raise ParserError(
            "python-docx is required to parse DOCX files; "
            "install it to enable uploads"
        ) from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ParserError(f"Failed to open DOCX: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    yield ParsedDocument(
        doc_id=str(uuid4()),
        text="\n".join(parts),
        source_filename=source_filename,
        section=None,
    )
